import docx


def style_heading(heading):
    doc = docx.Document()

    heading_paragraph = doc.add_paragraph(heading)
    heading_paragraph.style = doc.styles["Heading 1"]
    paragraph1 = doc.add_paragraph("Random text...")
    paragraph1.style.font.italic = True

    doc.add_paragraph("2nd Para...")
    
    doc.save('style_heading_test.docx')


style_heading("My Name")