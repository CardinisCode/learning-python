from datetime import datetime
import docx

class Recipient:
    def __init__(self, full_name, job_role, address_line):
        self.full_name = full_name
        self.job_role = job_role
        self.address_line = address_line
        self.current_date_obj = datetime.now().date()

    def get_recipient_details(self):
        details_string = self.full_name + "\n" + self.job_role + "\n"
        address_list = self.address_line.split(",")

        for line in address_list:
            details_string += line
        
        return details_string.lstrip()
    
    def get_open_greeting(self):
        return "Dear " + self.full_name + "," + "\n"

    def get_current_date_str(self):
        current_date = self.current_date_obj
        months_list = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
        month = int(current_date.strftime("%m"))
        month_string = months_list[month + 1]

        return current_date.strftime("%d") + " " + month_string + " " + current_date.strftime("%Y")


class JobHunter:
    def __init__(self, full_name, email_address, contact_number, postal_code):
        self.full_name = full_name
        self.email_address = email_address
        self.contact_number= contact_number
        self.postal_code = postal_code

    def get_job_hunters_details(self):
        string = self.email_address + " | " + self.contact_number + " | " + self.postal_code + "\n"
        return string


class CoverLetterBody:
    def __init__(self, first_paragraph, second_paragraph, third_paragraph, closing_paragraph):
        self.first_paragraph = first_paragraph
        self.second_paragraph = second_paragraph
        self.third_paragraph = third_paragraph
        self.closing_paragraph = closing_paragraph
        self.sign_off = "Kind regards,"


# Now lets try to use the above classes to put togeher an actual cover letter:
# Firstly we need to get input from the user:
def get_input_from_user():
    print("Lets start by getting your details.")
    users_full_name = input("Your Full name: ")
    users_email_address = input("Your email Address: ")
    users_contact_number = input("Your contact Number: ")
    users_postal_code = input("Your Postal Code: ")
    print()
    print("Now lets get details for the recipient of this cover letter.")
    recipients_full_name = input("Recipient's Full Name: ")
    recipients_job_title = input("Recipient's Job Title: ")
    recipients_address = input("Recipient's Address: ")
    print()
    print("The following questions will cover the content for the cover letter.")
    first_paragraph = input("The first paragraph of your cover letter: ")
    second_paragraph = input("The second paragraph of your cover letter: ")
    third_paragraph = input("The third paragraph of your cover letter (if applicable): ")
    closing_paragraph = input("The closing paragraph of your cover letter: ")
    print()

    input_details = {
        "users_full_name": users_full_name, 
        "users_email_address": users_email_address, 
        "users_contact_number": users_contact_number, 
        "users_postal_code": users_postal_code,
        "recipients_full_name": recipients_full_name, 
        "recipients_job_title": recipients_job_title, 
        "recipients_address": recipients_address, 
        "first_paragraph": first_paragraph, 
        "second_paragraph": second_paragraph, 
        "third_paragraph": third_paragraph,
        "closing_paragraph": closing_paragraph,
    }
    return input_details


def update_objects_with_user_input(user_input):
    # Lets start by instantiating our 3 objects:
    recipient = Recipient(
        user_input["recipients_full_name"], 
        user_input["recipients_job_title"], 
        user_input["recipients_address"]
    )
    job_hunter = JobHunter(
        user_input["users_full_name"], 
        user_input["users_email_address"], 
        user_input["users_contact_number"], 
        user_input["users_postal_code"]
    )
    cover_letter_body = CoverLetterBody(
        user_input["first_paragraph"], 
        user_input["second_paragraph"], 
        user_input["third_paragraph"],
        user_input["closing_paragraph"]
    )

    return recipient, job_hunter, cover_letter_body


def create_cover_letter():
    user_input = get_input_from_user()
    recipient, job_hunter, cover_letter_body = update_objects_with_user_input(user_input)

    doc = docx.Document()
    heading = doc.add_paragraph(job_hunter.full_name)
    heading.style = doc.styles["Heading 1"]
    doc.add_paragraph(job_hunter.get_job_hunters_details())
    doc.add_paragraph(recipient.get_recipient_details())
    doc.add_paragraph(recipient.get_current_date_str())

    doc.paragraphs[3].runs[0].add_break()
    doc.add_paragraph(recipient.get_open_greeting())
    doc.add_paragraph(cover_letter_body.first_paragraph)
    doc.paragraphs[5].runs[0].add_break()
    doc.add_paragraph(cover_letter_body.second_paragraph)
    doc.paragraphs[6].runs[0].add_break()
    doc.add_paragraph(cover_letter_body.third_paragraph)
    doc.paragraphs[7].runs[0].add_break()
    doc.add_paragraph(cover_letter_body.closing_paragraph)
    doc.paragraphs[8].runs[0].add_break()
    doc.add_paragraph(cover_letter_body.sign_off)

    doc.save('cover_letter_test.docx')


create_cover_letter()



