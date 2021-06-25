import docx


def write_word_doc(message, my_details):
    doc = docx.Document()
    doc.add_paragraph('My Name', 'Title')

    details_list = my_details.split(",")
    my_details = (" |".join(details_list))
    doc.add_paragraph(my_details)
    doc.paragraphs[1].runs[0].add_break()

    doc.add_paragraph(message)

    doc.save('new_doc.docx')

first_paragraph = ""
my_details = "my_email@gmail.com, 000, XX1 1XX"
write_word_doc("Random text...", my_details)