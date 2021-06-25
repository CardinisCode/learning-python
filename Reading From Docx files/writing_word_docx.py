import docx
from style_date_value import style_date_for_presentation
from prepare_strings import create_recipient_string
from prepare_strings import create_personal_info_string
from datetime import datetime, date



# Lets start by writing random text to a new file:
def write_to_new_docx(message):
    doc = docx.Document()
    doc.add_paragraph(message)
    doc.save('hello_world.docx')

# write_to_new_docx("Hello World!")

# Now to update the document
def update_docx(filename, message):
    doc = docx.Document(filename)
    doc.add_paragraph(message)
    doc.save(filename)

# update_docx('hello_world.docx', "Nice to meet you!")

# Now lets create a document & write multiple lines to it:
def create_docx_with_multiple_paragraphs(messages):
    doc = docx.Document()
    for para in messages:
        doc.add_paragraph(para)
    doc.save("multiple_paragraphs_updated.docx")

messages = [
    "Heading Title", 
    "First paragraph", 
    "Second Paragraph", 
    "Closing greeting"
]
# create_docx_with_multiple_paragraphs(messages)

# I want to tweak this concept a little to play around with an idea
def create_cv_docx(paragraph_lines):
    doc = docx.Document()
    for para in paragraph_lines:
        doc.add_paragraph(para)

    doc.save("write_cv_test.docx")

paragraph_lines = [
    "Hello ...", 
    "I am writing to express my interest in the position of...", 
    "Kind regards,"
]
create_cv_docx(paragraph_lines)

# I want to be able to customise the info that gets added to each paragraph. 
# So I'm going to try rather use a dict instead of a list.
# 
# Lets start by putting together the dict: 
paragraph_sections = {
    "my_name": "Andrea Folgado",
    "my_details": "ac.folgado@gmail.com | 073 7756 1277 \n",
    "todays_date": "",
    "recipients_address": "",
    "open_greeting": "", 
    "first_paragraph": "", 
    "second_paragraph": "", 
    "closing_paragraph": "", 
    "closing_greeting": "Kind regards,",
}

def create_cover_letter(paragraph_sections, recipients_name, recipients_address):
    doc = docx.Document()
    current_date = datetime.now().date()
    todays_date_str = style_date_for_presentation(current_date)
    paragraph_sections["todays_date"] = todays_date_str
    paragraph_sections["recipients_address"] = recipients_name + "\n" + recipients_address + "\n"
    paragraph_sections["open_greeting"] = "Dear " + recipients_name + "," + "\n"
    paragraph_sections["first_paragraph"] = "First paragraph" + "\n"
    paragraph_sections["second_paragraph"] = "Second Paragraph" + "\n"
    paragraph_sections["closing_paragraph"] = "Closing Paragraph" + "\n"

    for para in paragraph_sections.values():
        doc.add_paragraph(para)

    doc.save("write_coverletter_test.docx")

recipients_name = "Sherlock Holmes"
recipients_address = "21 Baker Street, London, W1U 8EQ"

# create_cover_letter(paragraph_sections, recipients_name)


def create_cover_letter_v2(recipients_name, recipients_address, recipients_job_description):
    doc = docx.Document()
    # Lets start by putting together the dict: 
    paragraph_sections = {
        "my_details": create_personal_info_string(),
        "address_string": "",
        "open_greeting": "", 
        "first_paragraph": "First paragraph", 
        "second_paragraph": "Second Paragraph", 
        "closing_paragraph": "Closing Paragraph", 
        "closing_greeting": "Kind regards,",
    }

    paragraph_sections["address_string"] = create_recipient_string(recipients_name, recipients_address, recipients_job_description)
    paragraph_sections["open_greeting"] = "Dear " + recipients_name + ","

    for para in paragraph_sections.values():
        doc.add_paragraph(para + "\n")

    doc.save("write_coverletter_test_v2.docx")

recipients_job_description = "Private Investigator"
create_cover_letter_v2(recipients_name, recipients_address, recipients_job_description)














